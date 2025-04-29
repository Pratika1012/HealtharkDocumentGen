import io
from docx import Document
import base64
import docx
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import pandas as pd
from docx.shared import Inches, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
import io
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_ALIGN_PARAGRAPH
import streamlit as st

def save_tables_to_excel(tables):
    output = io.BytesIO()
    # Step 2: Save each table as a sheet in an Excel file
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        for i, table in enumerate(tables):
            table.to_excel(writer, sheet_name=f'Table_{i+1}', index=False)

    output.seek(0)
    return output.getvalue()

def generate_excel_download_link(excel_data,name):
    # Encode the Excel file data to base64
    b64 = base64.b64encode(excel_data).decode()
    
    # Create the download link
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{name}.xlsx">Click here to download your Excel file</a>'
    return href



def set_table_border(table):
    # Loop through rows and cells to set borders for each cell
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')  # Create table cell borders element
            
            # Define border styles (top, bottom, left, right)
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Border type: single line
                border.set(qn('w:sz'), '4')       # Border size: 4 (1/8 pt)
                border.set(qn('w:space'), '0')    # No space between border and content
                border.set(qn('w:color'), '000000')  # Border color: black
                borders.append(border)  # Append border element
            
            tcPr.append(borders)  # Add borders to cell properties


def split_text_by_digits(text):
    # Regex pattern to match where a digit followed by a period and a space begins
    pattern = r"(?=\d.?\s)"  # Lookahead to detect digit followed by a space
    # Split the text into sections
    sections = re.split(pattern, text)
    # Remove empty strings and strip leading/trailing whitespace
    sections = [section.strip() for section in sections if section.strip()]
    return sections

def set_cell_border(cell, **kwargs):
    """
    Set cell border in a table.
    Args:
        cell: The table cell to modify.
        kwargs: Dictionary containing border attributes like sz, val, and color.
    """
    tc = cell._tc  # Access the XML representation of the cell
    tcPr = tc.get_or_add_tcPr()  # Get or create the <w:tcPr> element
    tcBorders = tcPr.find('w:tcBorders', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_element = tcBorders.find(f'w:{edge}', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if edge_element is None:
                edge_element = OxmlElement(f'w:{edge}')
                tcBorders.append(edge_element)
            # Set attributes for the edge
            edge_element.set('val', edge_data.get('val', 'single'))  # Border style
            edge_element.set('sz', str(edge_data.get('sz', '4')))  # Border thickness
            edge_element.set('space', '0')  # Space around the border
            edge_element.set('color', edge_data.get('color', 'auto'))  




# dynamic
def save_text_in_document_1(input,doc,flag,value=""):
    # print(text)
    
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    
    if flag==0:
        process_text_to_docx(doc, input)

    elif flag==1:
        num=value[0]
        image_extraction_paragraph = doc.add_paragraph(style='Normal')
        run = image_extraction_paragraph.add_run(f"{num}. Image Extraction")
        run.bold = True
        if input is not None:
        
            try:
                doc.add_picture(fr"ExtractedImages2\{input}", width=Inches(6.0))
            except:
                image_text = "Image is not found"
                process_text_to_docx(doc, image_text)


        else:
            image_text = "Image is not found"
            process_text_to_docx(doc, image_text)

    else:
        i = 0
        
        # ls0=split_text_by_digits(value)
        # st.write(ls0)
        
        for key,df in input.items():
            table_name=None
            try:
                heading=df[0]["heading"]
                table_name=df[0]["table_name"]
            except:
                heading=df[0]["heading"]

            table=df[0]["columns"]





            if len(table)>1:
                df0=pd.DataFrame(table)
            else:
                df0=pd.DataFrame(list(table[0].items()), columns=["Attribute", "Value"])

            


            if table_name is not None:

                    table_extraction = doc.add_paragraph(style='Normal')
                    run = table_extraction.add_run(f'{re.sub(r'\d+','',heading).strip()}')
                    run.bold = True

                    paragraph = doc.add_paragraph(f"Table: {table_name.strip()}")
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run1 = paragraph.runs[0]  # Access the first run in the paragraph
                    run1.bold = True
            else:
                paragraph = doc.add_paragraph(f"Table: {heading.strip()}")
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run1 = paragraph.runs[0]  # Access the first run in the paragraph
                run1.bold = True

            # ls=ls0[i].split(":")
            # print(ls)

            # if isinstance(df, list):
            #     if len(ls)>3:

            #         table_extraction = doc.add_paragraph(style='Normal')
            #         run = table_extraction.add_run(f'{re.sub(r'\d+','',ls[0]).strip()}')
            #         run.bold = True

            #         paragraph = doc.add_paragraph(f"Table: {ls[1].strip()}")
            #         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #         run1 = paragraph.runs[0]  # Access the first run in the paragraph
            #         run1.bold = True
            #     else:
            #         paragraph = doc.add_paragraph(f"Table: {ls[0].strip()}")
            #         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #         run1 = paragraph.runs[0]  # Access the first run in the paragraph
            #         run1.bold = True





            #     try:
            #         if len(df)>1:
            #             df0=pd.DataFrame(df)
            #         else:
            #             df0=pd.DataFrame(list(df[0].items()), columns=["Attribute", "Value"])
                        
                    

            #     except :
            #         df0=pd.DataFrame(list(df.items()), columns=["Attribute", "Value"])
            # elif  isinstance(df, dict):
            #         # if len(df) == 2 and all(isinstance(value, dict) for value in df.values()):
                    



            #         for key1,df1 in df.items():
            #             if isinstance(df1, (list, dict)):
            #                 if len(ls)>3:

            #                     table_extraction = doc.add_paragraph(style='Normal')
            #                     run = table_extraction.add_run(f'{re.sub(r'\d+','',ls[0]).strip()}')
            #                     run.bold = True

            #                     paragraph = doc.add_paragraph(f"Table: {ls[1].strip()}")
            #                     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #                     run1 = paragraph.runs[0]  # Access the first run in the paragraph
            #                     run1.bold = True
            #                 else:
            #                     paragraph = doc.add_paragraph(f"Table: {ls[0].strip()}")
            #                     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #                     run1 = paragraph.runs[0]  # Access the first run in the paragraph
            #                     run1.bold = True


            #                 if isinstance(df1, list):
            #                     try:
            #                         df0 = pd.DataFrame(df1)
            #                     except Exception as e:
            #                         print(f"Error converting sub-list to DataFrame: {e}")
            #                         df0 = pd.DataFrame(list(df1.items()), columns=["Attribute", "Value"])
            #                 elif isinstance(df1, dict):
            #                     df0 = pd.DataFrame(list(df1.items()), columns=["Attribute", "Value"])

                        

            #             else:
            #                 if len(ls)>3:

            #                     table_extraction = doc.add_paragraph(style='Normal')
            #                     run = table_extraction.add_run(f'{re.sub(r'\d+','',ls[0]).strip()}')
            #                     run.bold = True

            #                     paragraph = doc.add_paragraph(f"Table: {ls[1].strip()}")
            #                     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #                     run1 = paragraph.runs[0]  # Access the first run in the paragraph
            #                     run1.bold = True
            #                 else:
            #                     paragraph = doc.add_paragraph(f"Table: {ls[0].strip()}")
            #                     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #                     run1 = paragraph.runs[0]  # Access the first run in the paragraph
            #                     run1.bold = True

            #                 df0=pd.DataFrame(list(df.items()), columns=["Attribute", "Value"])
                            
            #                 break

                
            
                    


            if df0.empty:
                print("The DataFrame is empty. No table will be created.")
            else:
                # Create table in document

                if "Value" not in df0.columns:
                    first_col_values = df0.iloc[:, 0].unique()
                    second_col_values = df0.iloc[:, 1].unique()
                    if len(first_col_values) == 1 and len(second_col_values) == 1:
                        first_2_column= df0.iloc[:, :2]
                        df0= df0.iloc[:, 2:]



                        table = doc.add_table(rows=0, cols=len(first_2_column.columns))

                        row_cells = table.add_row().cells
                        row_cells[0].text = str(first_2_column.columns[0])
                        row_cells[0].paragraphs[0].runs[0].font.bold = True
                        row_cells[1].text = str("\n".join(first_2_column.iloc[:, 0].unique()))

                        row_cells = table.add_row().cells
                        row_cells[0].text = str(first_2_column.columns[1])
                        row_cells[0].paragraphs[0].runs[0].font.bold = True
                        row_cells[1].text = str("\n".join(first_2_column.iloc[:, 1].unique()))

                        set_table_border(table)


                
                    # Add rows from DataFrame
                    
                        

                    table = doc.add_table(rows=1, cols=len(df0.columns))

                    # Add column headers
                    hdr_cells = table.rows[0].cells
                    for j, col in enumerate(df0.columns):
                        if col is not None:  # Check for None values
                            hdr_cells[j].text = str(col)  # Ensure col is a string
                            hdr_cells[j].paragraphs[0].runs[0].font.bold = True
                            set_cell_background(hdr_cells[j], "F2F2F2")
                            

                    # Add rows from DataFrame
                    for index, row in df0.iterrows():
                        row_cells = table.add_row().cells
                        for j, cell in enumerate(row):
                            row_cells[j].text = str(cell)
                            
                    
                    first_col_values = df0.iloc[:, 0].unique()
                    second_col_values = df0.iloc[:, 1].unique()

                    # if len(first_col_values) == 1 and len(second_col_values) == 1:
                    #     # Merge the first column cells
                       
                       
                       
                    #     for row in table.rows[1:]:  # Skip header
                    #         row.cells[0].merge(table.rows[-1].cells[0])

                    #     # Merge the second column cells
                    #     for row in table.rows[1:]:  # Skip header
                    #         row.cells[1].merge(table.rows[-1].cells[1])

                    #     # Add vertical text for the merged cells
                    #     merged_cell_1 = table.rows[1].cells[0]  # First merged column
                    #     merged_cell_1.text = "\n".join(df0.iloc[:, 0].unique())
                    #     merged_cell_1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    #     set_cell_border(hdr_cells[j], top={"sz": 12, "val": "single", "color": "000000"},
                    #         bottom={"sz": 12, "val": "single", "color": "000000"},
                    #         left={"sz": 12, "val": "single", "color": "000000"},
                    #         right={"sz": 12, "val": "single", "color": "000000"})

                    #     merged_cell_2 = table.rows[1].cells[1]  # Second merged column
                    #     merged_cell_2.text = "\n".join(df0.iloc[:, 1].unique())
                    #     merged_cell_2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    #     set_cell_border(hdr_cells[j], top={"sz": 12, "val": "single", "color": "000000"},
                    #         bottom={"sz": 12, "val": "single", "color": "000000"},
                    #         left={"sz": 12, "val": "single", "color": "000000"},
                    #         right={"sz": 12, "val": "single", "color": "000000"})

                    #     for row in table.rows:
                    #         for cell in row.cells:
                    #             set_cell_border(cell, top={"sz": 6, "val": "single", "color": "000000"},
                    #                             bottom={"sz": 6, "val": "single", "color": "000000"},
                    #                             left={"sz": 6, "val": "single", "color": "000000"},
                    #                             right={"sz": 6, "val": "single", "color": "000000"})



                
                else:
                    table = doc.add_table(rows=0, cols=len(df0.columns))
                
                    # Add rows from DataFrame
                    for index, row in df0.iterrows():
                        row_cells = table.add_row().cells
                        for j, cell in enumerate(row):
                            row_cells[j].text = str(cell)

                set_table_border(table)
            doc.add_paragraph("")
            i += 1

def process_text_to_docx(doc,text):
    

    # Split the input text by lines
    lines = text.splitlines()

    for line in lines:
        # Handle headings (## for H2)
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        
        # Handle bold text (** for bold)
        elif '**' in line:
            p = doc.add_paragraph()
            bold_parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in bold_parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True  # Add bold text
                else:
                    p.add_run(part)  # Add regular text

        # Handle bullet points (* for bullet lists)
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='ListBullet')

        # Handle other normal paragraphs
        else:
            if line.strip() != "":  # Avoid adding empty lines
                doc.add_paragraph(line)

def set_cell_background(cell, color):
    """Set cell background shading color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Create a new shading element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)  # Set fill color
    tcPr.append(shd)


def save_text_in_document( text="", warning="", image_text="", image=None, tables=[], logo_path=None):
    # print(text)
    output = io.BytesIO()
    template_path = "templates/output_template.docx"
    # Load the template
    doc = Document(template_path)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    # Insert text
    process_text_to_docx(doc, text)

    process_text_to_docx(doc, warning)
   
    # Insert image
    table_extraction_paragraph = doc.add_paragraph(style='Normal')
    run = table_extraction_paragraph.add_run("Image Extraction")
    run.bold = True
    if image is not None:
        try:
            doc.add_picture(fr"ExtractedImages2\{image}", width=Inches(6.0))
        except:
            image_text = "Image is not found"
    
    process_text_to_docx(doc, image_text)

    # Add bold paragraph for Table Extraction
    table_extraction_paragraph = doc.add_paragraph(style='Normal')
    run = table_extraction_paragraph.add_run("Table Extraction")
    run.bold = True  # Make the text bold

    # Insert tables if available
    if tables is not None:
        i = 1
        for df in tables:
            table_extraction = doc.add_paragraph(style='Normal')
            run = table_extraction.add_run(f'Table-{i}')
            run.bold = True  # Make the text bold
            if df.empty:
                print("The DataFrame is empty. No table will be created.")
            else:
                # Create table in document
                table = doc.add_table(rows=1, cols=len(df.columns))

                # Add column headers
                hdr_cells = table.rows[0].cells
                for j, col in enumerate(df.columns):
                    if col is not None:  # Check for None values
                        hdr_cells[j].text = str(col)  # Ensure col is a string
                        hdr_cells[j].paragraphs[0].runs[0].font.bold = True
                        set_cell_background(hdr_cells[j], "F2F2F2")

                # Add rows from DataFrame
                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for j, cell in enumerate(row):
                        row_cells[j].text = str(cell)

                set_table_border(table)
            i += 1
    
    # Save the modified document to the output stream
    doc.save(output)

    # Reset the buffer position to the beginning
    output.seek(0)
    return output



def generate_word_download_link(doc_data, filename):
    b64 = base64.b64encode(doc_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx">Click here to download the Word document</a>'
    return href


def save_summaries_to_word(summaries):
    output = io.BytesIO()
    doc = Document()    
    process_text_to_docx(doc, summaries)
    doc.save(output)
    output.seek(0)  # Reset the buffer position to the beginning
    return output


# def save_text_in_document_cer(derived_tables=None,image=None, device_discription="",  tables=[],normal_text="",logo_path=None):
#     # print(text)
    

#     output = io.BytesIO()
#     template_path = "templates/output_template.docx"

#     # Load the template
#     doc = Document(template_path)
#     style = doc.styles["Normal"]
#     font = style.font
#     font.name = "Times New Roman"
    
    
#     # derived table

#     if derived_tables :
#         i = 1
#         for key,df in derived_tables.items():

#             #main heading of derived table
#             table_extraction = doc.add_paragraph(style='Normal')
#             run = table_extraction.add_run(f'{i} {key}')
#             run.bold = True
            

#             # table name
#             sub_key=list(df.keys()) 
#             paragraph = doc.add_paragraph(f"Table: {sub_key[0]}")
#             # Center align the paragraph
#             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             # Make the text bold
#             run = paragraph.runs[0]  # Access the first run in the paragraph
#             run.bold = True

#             try:
#                 df=pd.DataFrame(df[sub_key[0]])
                

#             except :
#                 df=pd.DataFrame(list(df[sub_key[0]].items()), columns=["Attribute", "Value"])



#             if df.empty:
#                 print("The DataFrame is empty. No table will be created.")
#             else:
#                 # Create table in document

#                 if "Value" not in df.columns:
#                     table = doc.add_table(rows=1, cols=len(df.columns))

#                     # Add column headers
#                     hdr_cells = table.rows[0].cells
#                     for j, col in enumerate(df.columns):
#                         if col is not None:  # Check for None values
#                             hdr_cells[j].text = str(col)  # Ensure col is a string
#                             hdr_cells[j].paragraphs[0].runs[0].font.bold = True
#                             set_cell_background(hdr_cells[j], "F2F2F2")

#                     # Add rows from DataFrame
#                     for index, row in df.iterrows():
#                         row_cells = table.add_row().cells
#                         for j, cell in enumerate(row):
#                             row_cells[j].text = str(cell)
                
#                 else:
#                     table = doc.add_table(rows=0, cols=len(df.columns))
                
#                     # Add rows from DataFrame
#                     for index, row in df.iterrows():
#                         row_cells = table.add_row().cells
#                         for j, cell in enumerate(row):
#                             row_cells[j].text = str(cell)

#                 set_table_border(table)
#             doc.add_paragraph("")
#             i += 1
    
    
#     # Insert text
#     process_text_to_docx(doc, device_discription)



#     # Insert image
#     image_extraction_paragraph = doc.add_paragraph(style='Normal')
#     run = image_extraction_paragraph.add_run("Image Extraction")
#     run.bold = True
#     if image is not None:
        
#         try:
#             doc.add_picture(fr"ExtractedImages2\{image}", width=Inches(6.0))
#         except:
#             image_text = "Image is not found"
#             process_text_to_docx(doc, image_text)


#     else:
#         image_text = "Image is not found"
#         process_text_to_docx(doc, image_text)

    


#     table_extraction_paragraph = doc.add_paragraph(style='Normal')
#     run = table_extraction_paragraph.add_run("Table Extraction")
#     run.bold = True  # Make the text bold



#     # Insert tables if available
#     if tables is not None:
#         i = 1
#         for df in tables:
#             table_extraction = doc.add_paragraph(style='Normal')
#             run = table_extraction.add_run(f'Table-{i}')
#             run.bold = True  # Make the text bold
#             if df.empty:
#                 print("The DataFrame is empty. No table will be created.")
#             else:

                
#                 # Create table in document
#                 table = doc.add_table(rows=1, cols=len(df.columns))

#                 # Add column headers
#                 hdr_cells = table.rows[0].cells
#                 for j, col in enumerate(df.columns):
#                     if col is not None:  # Check for None values
#                         hdr_cells[j].text = str(col)  # Ensure col is a string
#                         hdr_cells[j].paragraphs[0].runs[0].font.bold = True
#                         set_cell_background(hdr_cells[j], "F2F2F2")

#                 # Add rows from DataFrame
#                 for index, row in df.iterrows():
#                     row_cells = table.add_row().cells
#                     for j, cell in enumerate(row):
#                         row_cells[j].text = str(cell)

                  

#                 set_table_border(table)
#             i += 1

#     process_text_to_docx(doc, normal_text)




#     # Add bold paragraph for Table Extraction
    
    
#     # Save the modified document to the output stream
#     doc.save(output)

#     # Reset the buffer position to the beginning
#     output.seek(0)
#     return output