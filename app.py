import streamlit as st
import os
import io
import logging
import time
import shutil
from Features.Extraction_module import extraction
from Features.Summarization import summarized_Document
# from Features.scrap import scrapping
from Features.scrap import scrapping
from doc_generate import generate_word_download_link
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
# from Features.scrap import scrapping


# Reset logging configuration
for handler in logging.root.handlers[:]:
    logging.root.removeHandler(handler)

logger = logging.getLogger()

# Check if handlers are already present
if not logger.hasHandlers():
    # Set logging level
    logger.setLevel(logging.INFO)

    # Create handlers
    file_handler = logging.FileHandler('logs/app.log')
    console_handler = logging.StreamHandler()

    # Set levels for handlers
    file_handler.setLevel(logging.INFO)
    console_handler.setLevel(logging.INFO)

    # Create formatter and add to handlers
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)

    # Add handlers to logger
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
# logging.basicConfig(filename='app.log', level=logging.INFO, filemode='a',
#                     format='%(asctime)s - %(levelname)s - %(message)s')
logging.info("App started")
def add_logo_to_docx(doc_data, logo_path="pdf_logo.png", text="The world leader in serving science"):
    """Adds a logo to the left side of the top header section and specified text to the right side.

    Args:
        doc_data (bytes): The content of the Word document as a byte stream.
        logo_path (str, optional): The path to the logo image file. Defaults to "pdf_logo.png".
        text (str, optional): The text to add to the header. Defaults to "world leader in serving science".

    Returns:
        bytes: The modified Word document content as a byte stream.
    """
    
    document = Document(doc_data)

    # Get the first section's header
    section = document.sections[0]
    header = section.header

    # Create a table with two cells: one for the logo, one for the text
    table = header.add_table(rows=1, cols=2, width=Inches(6))
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align table to left

    # Adjust column widths
    table.columns[0].width = Inches(1.5)  # Width for the logo
    table.columns[1].width = Inches(4.5)  # Width for the text

    # Add logo to the first cell
    logo_cell = table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.25))

    # Add text to the second cell, aligning it to the right
    text_cell = table.cell(0, 1)
    text_paragraph = text_cell.paragraphs[0]
    text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Create a run with a manual line break
    bold_run = text_paragraph.add_run("The world leader\nin serving science")
    bold_run.font.name = "Times New Roman"
    bold_run.bold = True  # Make the text bold

    # Optionally set margins for all sections
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Save the document and return as byte stream
    document.save("temp.docx")
    with open("temp.docx", "rb") as f:
        return f.read()
    
# Title of the app
st.markdown(
    "<h1 style='text-align: center; color: #2F5C87;'>Document Generator</h1>",
    unsafe_allow_html=True
)
st.markdown("""
    <style>
    /* Styling for the main submit button */
    div.stButton > button {
        background-color: #2F5C87;  /* Dark blue background */
        width: 200px;           /* Adjust width */
        height: 40px;           /* Adjust height */
        border: none;           /* Remove border */
        cursor: pointer;        /* Change cursor on hover */
        border-radius: 25px;    /* Rounded corners */
        color: White;           /* White text color */
        border: 2px solid white; /* Border */
        margin-top: 5px;
    }

    /* Hover effect for the buttons */
    div.stButton > button:hover {
        background-color: white;
        color: #2F5C87;
        border: 2px solid #2F5C87;
    }
        .sidebar-button.clicked {
        background-color: black !important;
        color: white !important;
    }
    </style>
""", unsafe_allow_html=True)

image_name=[]
# Example: Add your logo (local or online image path)
logo_url = "Utilites/logo1.jpg"  # Replace with your logo URL or local path

# Add the logo to the sidebar
st.sidebar.image(logo_url, width=200, use_column_width=False, caption=None )

# Initialize session state variables if they don't exist
if 'selected_feature' not in st.session_state:
    st.session_state.selected_feature = None
    
col1, col2, col3 = st.columns(3)

st.sidebar.header("Document Generator Menu")
with st.sidebar:
    if st.button("Data Extraction"):
        st.session_state.selected_feature = 'generator'

    if st.button("Doc Summarization"):
        st.session_state.selected_feature = 'summarization'

    if st.button("Web Scraping"):
        st.session_state.selected_feature = 'scraping'

# Logic to display content based on button selection
if st.session_state.selected_feature == 'generator':
    st.subheader("📄 Data Extraction")
    st.write("")

    input_file_text = st.file_uploader(
        "🔽 Upload Document for Extraction ",
        accept_multiple_files=True,
        type=['pdf', 'docx']
    )
    
    if input_file_text:
        st.success(f"{len(input_file_text)} file(s) uploaded successfully!")
        
    st.write("")
    st.write("")
    
    # Now show the Output File Uploader
    option = st.radio(
        "Choose Type of extraction",
        ("CER Building","Normal Extraction")
    )
    
    # Display result based on the selected option
    flag=0
    indication=0
     
 
    if option == "CER Building":
        sub_option = st.radio(
        "Choose type of approach",
        ("Similar Device","Equivalent Device")
        )

        if sub_option=="Similar Device":
            indication=0
        else:
            indication=1

        logging.info("Option 2 selected")

        output_file_text = "templates\CER Refernce(05-12-24).docx"
        model_number=""
        flag=1
        
        
    elif option == "Normal Extraction":
        st.write("Please enter the model number series for which you would like to generate output")
        text_to_extract = st.text_input("", placeholder="Type model number series and then press enter")
        model_number = text_to_extract
        logging.info("Option 1 selected")
        logging.info(model_number)
        output_file_text = "templates/Task1_Data_Extraction_Template_2 (1) (1).docx"
        flag=0
    st.write("")
    st.write("")
    token_info=[] 
    input_tokens = 0
    output_tokens = 0
    total_tokens = 0  
 
    if st.button("Submit") :
            start_time = time.time()
            text,documnet_name,token_info,input_tokens,output_tokens,total_tokens  = extraction(flag,indication,model_number,input_file_text,output_file_text)
            response_time = time.time() - start_time

            logging.info(f"Total Tokens for Doc Generate------------------------------------------")
            logging.info(f"Input Tokens: {input_tokens}")
            logging.info(f"Output Tokens: {output_tokens}")
            logging.info(f"Total Tokens: {total_tokens}")
            logging.info(f"Response generation time: {response_time:.2f} seconds")

            # file_name = os.path.splitext(documnet_name)[0]
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
           
            if flag==0:
                prefix = "Regulatory_Document" 
                final_file_name = f"{prefix}_{documnet_name}_{timestamp}"

                if text: 
                    st.subheader("📥 Download generated file")
                    st.subheader("Token Information")
                    st.write(f"Input Tokens: {input_tokens}")
                    st.write(f"Output Tokens: {output_tokens}")
                    st.write(f"Total Tokens: {total_tokens}")
                    st.write(f"Response generation time: {response_time:.2f} seconds")
                else:  
                    st.write("No text available for download.")
                    st.button("Download", disabled=True)
                st.markdown(generate_word_download_link(text.getvalue(),final_file_name), unsafe_allow_html=True)
            
            else:
                prefix="CER" 
                final_file_name = f"{prefix}_{documnet_name}_{timestamp}"
             
            
                if text: 
                    st.subheader("📥 Download generated file")
                    st.subheader("Token Information")
                    st.write(f"Input Tokens: {input_tokens}")
                    st.write(f"Output Tokens: {output_tokens}")
                    st.write(f"Total Tokens: {total_tokens}")
                    st.write(f"Response generation time: {response_time:.2f} seconds")
                else:  
                    st.write("No text available for download.")
                    st.button("Download", disabled=True)
                st.markdown(generate_word_download_link(text.getvalue(),final_file_name), unsafe_allow_html=True)
        
#doc summarization new (12/11/2024)
if st.session_state.selected_feature == 'summarization':
    st.subheader("📝 Document Summarization")

    input_files = st.file_uploader("🔽 Upload Document for Summarization",
                                   accept_multiple_files=True,
                                   type=['pdf', 'docx'])

    if input_files:
        st.success(f"{len(input_files)} file(s) uploaded successfully!")

    # Load the reference template
    template_path = os.path.join("templates", "Summarization_template.docx")
    
    try:
        # Reading the template file
        template_doc = Document(template_path)
        template_text = "\n".join([para.text for para in template_doc.paragraphs])
    except FileNotFoundError:
        st.error("The reference template file was not found. Please check the path and try again.")
     
        st.stop()

    if st.button("Generate Summary"):
        start_time = time.time()
        summary_doc, total_tokens = summarized_Document(input_files, template_path)

        response_time = time.time() - start_time
        logging.info(f"Total Tokens: {total_tokens}")
        logging.info(f"Response generation time: {response_time:.2f} seconds")

        if summary_doc:
            st.subheader("📥 Download generated file")
            st.subheader("Token Information")

            st.write(f"Input Tokens: {total_tokens.get('input_tokens', 0)}")
            st.write(f"Output Tokens: {total_tokens.get('output_tokens', 0)}")
            st.write(f"Total Tokens: {total_tokens.get('total_tokens', 0)}")

            st.write(f"Response generation time: {response_time:.2f} seconds")
            summary_doc = io.BytesIO(summary_doc)
            st.markdown(generate_word_download_link(summary_doc.getvalue(), "Summary_Documents"), unsafe_allow_html=True)
        else:
            st.write("No text available for download.")
            st.button("Download", disabled=True)
    else:
        st.warning("Please upload at least one document before generating the summary.")


# Streamlit application
if st.session_state.selected_feature == 'scraping':
    st.subheader("🌐 Document Scraping")
    thermoDevice = st.text_input("Enter the name of the Thermo Scientific device:")
    previousGenDevice = st.text_input("Enter the name of the previous generation device:")
    user_keyword = st.text_input("Enter the keyword to search device:")
    refined_keyword = st.text_input("Enter Refined Keyword:")

    if thermoDevice and previousGenDevice:
        st.write(f"Scraping content from: \nThermo Scientific device: {thermoDevice}\nPrevious generation device: {previousGenDevice}")
    
    technical_data_sheet = st.file_uploader("🔽 Upload Technical Data Sheet for Device to be analyzed",
                                   accept_multiple_files=True,
                                   type=['pdf', 'docx'])
    prev_gen_data_sheet = st.file_uploader("🔽 Technical Data Sheet for Previous Generation Device to be analyzed",
                                   accept_multiple_files=True,
                                   type=['pdf', 'docx'])
    if st.button("Scrape"):
        start_time = time.time()
        byte_stream, token_info = scrapping(thermoDevice, previousGenDevice, user_keyword, refined_keyword, technical_data_sheet, prev_gen_data_sheet)
        response_time = time.time() - start_time
        logging.info(f"Total Tokens: {token_info}")
        logging.info(f"Response generation time: {response_time:.2f} seconds")

        if byte_stream:
            st.subheader("📥 Download generated file")
            st.subheader("Token Information")

            st.write(f"Input Tokens: {token_info.get('input_tokens', 0)}")
            st.write(f"Output Tokens: {token_info.get('output_tokens', 0)}")
            st.write(f"Total Tokens: {token_info.get('total_tokens', 0)}")

            st.write(f"Response generation time: {response_time:.2f} seconds")
            # summary_doc = io.BytesIO(byte_stream)           
            st.markdown(generate_word_download_link(byte_stream.getvalue(), "web_scrapping"), unsafe_allow_html=True)
        else:
            st.write("No text available for download.")
            st.button("Download", disabled=True)