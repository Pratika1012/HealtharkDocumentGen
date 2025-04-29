from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
import time
import openai
from bs4 import BeautifulSoup
import json
import pandas as pd
import streamlit as st
import logging
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time, os
from selenium.webdriver.chrome.service import Service
# WebDriver Manager for ChromeDriver
from webdriver_manager.chrome import ChromeDriverManager
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import fitz
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment
from openpyxl.utils.dataframe import dataframe_to_rows
# Load the configuration file
with open('Config/configuration.json', 'r') as f:
    config = json.load(f)

# Set the OpenAI API key
openai.api_key = config['api_key']
all_tokens = []
def md_device(button_id, filename, download_dir):
    """
    Downloads medical device data and renames the file.

    Args:
        button_id (str): The ID of the button to click.
        filename (str): The desired filename for the downloaded file.
        download_dir (str): The directory where the downloaded file will be saved.

    Returns:
        None
    """
    download_dir = os.path.abspath(download_dir)
    os.makedirs(download_dir, exist_ok=True)
    print("download dir:", download_dir)
    # Set Chrome options to specify download directory
    options = webdriver.ChromeOptions()
    prefs = {"download.default_directory": download_dir}
    options.add_experimental_option("prefs", prefs)
    options.add_argument("--headless")

    # Create Chrome webdriver object with options
    driver = webdriver.Chrome(options=options)

    try:
        # Open the URL
        driver.get("https://cdscomdonline.gov.in/NewMedDev/ListOfApprovedDevices")

        # Wait for the specified button to be clickable
        wait = WebDriverWait(driver, 20)
        button = wait.until(EC.element_to_be_clickable((By.ID, button_id)))

        # Click the button
        button.click()

        # Wait for the data to load (adjust the time if necessary)
        time.sleep(120)

        # Find the Excel button and click it
        excel_button = driver.find_element(By.CSS_SELECTOR, ".dt-button.buttons-excel.buttons-html5")
        excel_button.click()

        # Wait for the download to complete (adjust the time if necessary)
        time.sleep(10)

        # Rename the downloaded file
        downloaded_file = os.path.join(download_dir, "MD - Medical Devices.xlsx")
        renamed_file = os.path.join(download_dir, filename)

        # Ensure the downloaded file exists before renaming
        if os.path.exists(downloaded_file):
            if os.path.exists(renamed_file):
                os.remove(renamed_file)  # Delete the existing file
            os.rename(downloaded_file, renamed_file)
        else:
            print(f"File {downloaded_file} not found for renaming.")
        return renamed_file
    finally:
        # Close the browser regardless of exceptions
        driver.quit()

def scrap_process():
    # Specify the download directory
    download_dir = r"Features"
    print("entered scrap_process")
    # Download and rename files
    import_file =  md_device(button_id="acton2", filename="import_data.xlsx", download_dir=download_dir)
    print(import_file)
    manufacture_file = md_device(button_id="acton1", filename="manufactur_data.xlsx", download_dir=download_dir)

    if import_file and manufacture_file:
            # Read both Excel files, dropping the first row
            import_df = pd.read_excel(import_file, skiprows=1)
            manufacture_df = pd.read_excel(manufacture_file, skiprows=1)
            # Combine the dataframes
            combined_df = pd.concat([import_df, manufacture_df], ignore_index=True)
            combined_df['S No.'] = range(1, len(combined_df) + 1)
            # Save the combined dataframe to a new Excel file
            combined_file_path = os.path.join(os.getcwd(),"Features", "MD - Medical Devices (3).xlsx")
            combined_df.to_excel(combined_file_path, index=False)
            print(f"Combined file saved as {combined_file_path}")
    else:
        print("Failed to download one or both of the Excel files")

def initialize_driver():
    """Initialize the Selenium WebDriver with required options."""
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1920,1080")
    driver = webdriver.Chrome(options=chrome_options)
    return driver

def scrape_device_details(driver, competitor):
    """
    Scrape details for a specific device from Google search results.
    """
    device_name = competitor
    driver.get("https://www.google.com")
    search_box = driver.find_element(By.NAME, "q")
    search_box.clear()
    search_box.send_keys(device_name)
    search_box.send_keys(Keys.RETURN)
    time.sleep(5)

    results_data = []
    try:
        search_results = driver.find_elements(By.CSS_SELECTOR, "div.yuRUbf a")
        for i in range(2):
            if i >= len(search_results):
                break
            search_results[i].click()
            time.sleep(5)
            soup = BeautifulSoup(driver.page_source, "html.parser")
            page_title = soup.title.string if soup.title else "No Title"
            page_text = soup.get_text(separator="\n", strip=True)
            results_data.append({
                "Result Number": i + 1,
                "URL": driver.current_url,
                "Title": page_title,
                "Content": page_text
            })
            driver.back()
            time.sleep(5)
            search_results = driver.find_elements(By.CSS_SELECTOR, "div.yuRUbf a")
    except Exception as e:
        print("An error occurred while extracting search results:", e)
    return results_data

def format_scraped_data(results):
    """
    Format scraped data into a single string for input to the analysis function.
    """
    return "\n".join(
        f"Result Number: {item['Result Number']}\n"
        f"URL: {item['URL']}\n"
        f"Title: {item['Title']}\n"
        f"Content: {item['Content']}\n"
        for item in results
    )



def analyze_google_data_tableA(thermoDevice, previousDevice, googleData_thermoDevice, googleData_previousDevice):
    """
    Analyze and compare data for the current and previous generation devices using OpenAI API.
    """
    
    # prompt = f"""
    # Here the current device is: {thermoDevice} and previous generation device is: {previousDevice}.
    # Now, From the Current Device extract: {googleData_thermoDevice} and previous generation device extract: {googleData_previousDevice} 
    # Extract the information for the following key "parameters" for both the devices, try to search exact "parameters" first 
    # then if the parameter word is not found then try to find information from full content:
    # "Temperature Range","Electrical Requirements","Power Plug / Power Cord","Application Environment","Refrigeration System","Condenser Type","Expansion Device",
    # "Evaporator Type","Defrost Method","Control Sensor","Connectivity / Remote Outputs","Adjustable Warm/Cold Alarms","Controller Level","Compressor Safeguard",
    # "Amperage","Power Switch","Controller Type","Refrigerant","Door Seal"

    # I want you to give output in JSON only.
    # Main key of that JSON has to be the "parameters" followed by extracted values for previous generation device and current device and finally a flag of 'Similar' or 'Different' 
    # depending on the comparision of the extracted values.
    # """
    prompt = f"""
    Here the current device is: {thermoDevice} and previous generation device is: {previousDevice}.
    Now, From the Current Device extract: {googleData_thermoDevice} and previous generation device extract: {googleData_previousDevice}
    Extract the information for the following key "parameters" for both the devices, try to search exact "parameters" first and find it's corresponding value,
    if the parameter word is not found then try to find information from full content:
    "Temperature Range","Electrical Requirements","Power Plug / Power Cord","Application Environment","Refrigeration System","Condenser Type","Expansion Device",
    "Evaporator Type","Defrost Method","Control Sensor","Connectivity / Remote Outputs","Adjustable Warm/Cold Alarms","Controller Level","Compressor Safeguard",
    "Instrument Rated Current","Power Switch","Controller Type","Refrigerant","Door Seal".

    I want you to give output in JSON only.
    Main key of that JSON has to be the "parameters" followed by extracted values for previous generation device and current device
    and finally a flag of 'Similar' if values are same or even close to one another and flag 'Different' if they are very different from each other.
    For example: 208-230V is Similar to 230V."""

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an expert in competitor product analysis."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.1
    )
    # response_time = time.time() - start_time
    start_time = time.time()
    # Return the generated content if available
    if response and response['choices']:
        input_tokens = response['usage']['prompt_tokens']
        output_tokens = response['usage']['completion_tokens']
        total_tokens = response['usage']['total_tokens']
        all_tokens.append({
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': total_tokens,
            'response_time': time.time() - start_time
        })


        # # Log the token usage and response time
        # logging.info("process_scrapping_with_GPT....")
        # logging.info(f"Input Tokens: {input_tokens}")
        # logging.info(f"Output Tokens: {output_tokens}")
        # logging.info(f"Total Tokens: {total_tokens}")
        # logging.info(f"Response generation time: {response_time:.2f} seconds")

    #     return response['choices'][0]['message']['content'], {
    #         'input_tokens': input_tokens,
    #         'output_tokens': output_tokens,
    #         'total_tokens': total_tokens,
    #         'response_time': response_time
    #     }
    # else:
    #     return None, None
    return response['choices'][0]['message']['content']

def create_competitor_list_TableC(tableB_df, thermo_device):
    """
    Generate a list of competitor devices from the refined DataFrame.
    """
    competitors_tableC = tableB_df.apply(lambda row: f"{row['Company Name']} {row['Product Name']}", axis=1).tolist()
    competitors_tableC[0]= thermo_device

    return competitors_tableC


def clean_json_response(response):
    """
    Clean the JSON response from OpenAI API by removing markdown formatting.
    """
    if response.startswith("```json") and response.endswith("```"):
        response = response[7:-3].strip()
    return response

def create_dataframe_from_analysis(analysis, previousDevice, currentDevice):
    """
    Convert analysis JSON into a structured DataFrame.
    """
    print("Type of analysis:", type(analysis))
    print("Content of analysis:", repr(analysis))
        # Remove Markdown-style code block markers
    if analysis.startswith("```json"):
        analysis = analysis[7:]  # Remove the leading ```json
    if analysis.endswith("```"):
        analysis = analysis[:-3]  # Remove the trailing ```
        # Parse JSON
    cleaned_analysis = analysis.strip()  # 
    if isinstance(cleaned_analysis, dict):
        data = cleaned_analysis  # Already parsed
    else:
        data = json.loads(cleaned_analysis)

    tableA_df = pd.DataFrame.from_dict(data, orient='index')
    tableA_df.rename(
        columns={
            "Previous Generation Device": previousDevice,
            "Current Device": currentDevice,
            "Comparison": "Comparison"
        },
        inplace=True
    )
    tableA_df.reset_index(inplace=True)
    tableA_df.rename(columns={"index": "Parameters"}, inplace=True)
    return tableA_df


def load_excel_data(file_path, sheet_name):
    """
    Load data from an Excel file into a DataFrame.
    """
    return pd.read_excel(file_path, sheet_name=sheet_name)

def filter_rows_by_keyword(df, keyword):
    """
    Filter rows in the DataFrame that contain the given keyword in any column.
    """
    filtered_rows = df[df.apply(lambda row: row.astype(str).str.contains(keyword, case=False, na=False).any(), axis=1)]
    return filtered_rows.reset_index()

def call_openai_api(prompt, model="gpt-4o"):
    """
    Call the OpenAI API to process a given prompt and return the response.
    """
    messages = [{"role": "user", "content": prompt}]
    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=messages,
            temperature=0
        )
        start_time = time.time()
        # Return the generated content if available
        if response and response['choices']:
            input_tokens = response['usage']['prompt_tokens']
            output_tokens = response['usage']['completion_tokens']
            total_tokens = response['usage']['total_tokens']
            all_tokens.append({
                'input_tokens': input_tokens,
                'output_tokens': output_tokens,
                'total_tokens': total_tokens,
                'response_time': time.time() - start_time
            })
        return response.choices[0].message["content"]
    except Exception as e:
        print(f"OpenAI API call failed: {e}")
        return None

def extract_json_from_response(response):
    """
    Extract JSON content from the OpenAI response using regex.
    """
    try:
        json_match = re.search(r"\[.*\]", response, re.DOTALL)  # Extract JSON content
        if json_match:
            json_data = json_match.group(0)  # Get the matched JSON string
            return json.loads(json_data)  # Convert JSON string to a Python object
    except Exception as e:
        print(f"Error extracting JSON from response: {e}")
    return None

def refine_data(filtered_rows, refined_keyword):
    """
    Refine the filtered DataFrame further using the OpenAI API based on a refined keyword.
    """
    dataframe_string = filtered_rows.to_string(index=False)
    prompt = f"""You are an assistant helping with data refinement tasks. Here is the entire DataFrame filtered from the original data:{dataframe_string}
    The user wants to refine the search using the keyword: "{refined_keyword}". 
    Include rows that have any column similar to the description mentioned in the keyword. Try and include as many rows as possible. 
    
    Provide the output as a JSON list of dictionaries where each dictionary represents a row that matches this refined search criteria. 
    Give the json list of dictionary only as an output, nothing else."""
    response = call_openai_api(prompt)
    if response:
        return extract_json_from_response(response)
    return None

def create_competitor_list(refined_df, thermo_device):
    """
    Generate a list of competitor devices from the refined DataFrame.
    """
    competitors = refined_df.apply(lambda row: f"{row['Name of Device and Device Class']} {row['Brand Name']}", axis=1).tolist()
    competitors.insert(0, thermo_device)
    return competitors
def analyze_google_data_TableB(data, competitor):
    prompt = f"""
    Analyze the following search results about {competitor}:
    {data}
    
    Just a general guide: If the device series belongs to Thermo Fisher, Focus on device models ending with "LV" for Product Name.

    Extract and organize the details into a JSON object with the following keys:
    - Company Name: (Manufacturer, issuer, etc. The company which manufactures the device)
    - Category: Product Category of the device. For example: -86°C ULT Freezers, -40°C ULT Freezers, Precriptive Bandages, etc.
    - Product Name:(Various model numbers of the device, seperate them by comma if multiple model numbers or generic of the series)
    - Description: (Generic elaborated description of a model series without specific details that can vary across the model numbers. Limit of description is 80 words)
    **If any of the above keys is not available in the data and 
    you are also not certain and aware of that information then please fill with "Data not found"
    Return only in JSON format with mentioned Keys only.
    """
    # start_time = time.time()
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an expert in competitor product analysis."},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.0
        )
    start_time = time.time()
        # Return the generated content if available
    if response and response['choices']:
        input_tokens = response['usage']['prompt_tokens']
        output_tokens = response['usage']['completion_tokens']
        total_tokens = response['usage']['total_tokens']
        all_tokens.append({
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': total_tokens,
            'response_time': time.time() - start_time
        })   

    # response_time = time.time() - start_time
    return response['choices'][0]['message']['content']

competitor_analysis = []
def scrap_process_tableB_df(file_path, sheet_name,thermoDevice, previousGenDevice, user_keyword, refined_keyword):
    thermo_device = thermoDevice
    df = load_excel_data(file_path, sheet_name)
    filtered_rows = filter_rows_by_keyword(df, user_keyword)

    # Refine data using OpenAI API
    refined_data = refine_data(filtered_rows, refined_keyword)
    competitors = []
    if refined_data:
        refined_df = pd.DataFrame(refined_data)
        print("\nRefined DataFrame:")
        print(refined_df)

        # Generate competitor list
        competitors = create_competitor_list(refined_df, thermo_device)
        print("\nCompetitor List:")
        print(competitors)
    else:
        print("Failed to refine data using OpenAI.")
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1920,1080")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    # driver.get("https://www.google.com")

    for competitor in competitors:
        try:
            print(f"Searching for: {competitor}")
            print(type(competitor))
            google_data = scrape_device_details(driver,competitor)
            print(f"Found {type(google_data)}")
            print(google_data)
            google_data_string = "\n".join(
                f"Result Number: {item['Result Number']}\n"
                f"URL: {item['URL']}\n"
                f"Title: {item['Title']}\n"
                f"Content: {item['Content']}\n"
                for item in google_data
            )
            print(google_data_string)
            analysis = analyze_google_data_TableB(google_data_string, competitor)
            print(analysis)
            competitor_analysis.append({"Competitor": competitor, "Analysis": analysis})
        except Exception as e:
            print(f"Error during analysis for competitor '{competitor}': {e}")

    processed_data=[]

    for item in competitor_analysis:
        analysis = item.get('Analysis', '')
        print(type(analysis))
        print(analysis)
        print("***********************************************")
            # If 'analysis' is a tuple, extract the first element (the JSON string)
        # if isinstance(analysis, tuple):
        #     json_string = analysis[0]  # Extract the JSON string
        # else:
        #     json_string = analysis  # Assume it's already a string

        # Try to parse the Analysis JSON; if it fails, skip this row
        try:
            analysis_data = json.loads(analysis.strip())
        except json.JSONDecodeError:
            print(f"Invalid JSON for analysis: {analysis}")
            continue

        # Add the parsed analysis to the processed data
        processed_data.append(analysis_data)

        # Convert the processed data into a DataFrame
    tableB_df = pd.DataFrame(processed_data)

    # Perform any necessary trimming or cleanup
    tableB_df = tableB_df.applymap(lambda x: x.strip() if isinstance(x, str) else x)
    print(tableB_df)
    return  tableB_df

def analyze_google_data_TableC(data, competitor):
    prompt = f"""
    Analyze the following search results about {competitor}:
    {data}

    Just a general guide: 
    If the device series belongs to Thermo Fisher, Focus on device models ending with "LV" for model information and related values
    like Storage Volume, Voltage, etc. Dont miss any such device ending with "LV".

    Extract and organize the details into a JSON object with the following keys:
    - License Holder: (Manufacturer, issuer, etc. The company which manufactures the device)
    - Description: (Generic description of device for a model series without specific details that can vary across the model numbers)
    - Intended Use: The Intended use of the Device series in about 70 words.
    - Storage Volume: (Capacity of the device for each model with the model number, seperate them by comma if multiple model numbers)
    - Temperature Range: (Temperature range of the device for each model with the model number, seperate them by comma if multiple model numbers)
    - Orientation: (What is the orientation of device. for example: Upright)
    - Voltage: (Voltage of the device for each model with the model number, seperate them by comma if multiple model numbers)
    - Material: (Type of material used in the device)
    - Insulation: (Type of insulation)
    - Refrigeration system: (What kind of refigeration system is used)    
    - Alarms: (Type of alarms system)
    - Indoor/Outdoor Usage: (Mention if device can be used Indoor or both Indoor and Outdoor)   

    **If any of the above keys is not available in the data and 
    you are also not certain and aware of that information then please fill with "Data not found"
    Return only in JSON format with mentioned Keys only.
    """

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an expert in competitor product analysis."},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.0
    )
    start_time = time.time()
    # Return the generated content if available
    if response and response['choices']:
        input_tokens = response['usage']['prompt_tokens']
        output_tokens = response['usage']['completion_tokens']
        total_tokens = response['usage']['total_tokens']
        all_tokens.append({
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': total_tokens,
            'response_time': time.time() - start_time
        })
    return response['choices'][0]['message']['content']
def process_and_transpose_data(tableC_analysis):
    # Create an empty list to hold processed rows
    processed_data_tableC = []
    competitors = []

    # Iterate through each entry in the tableC_analysis list
    for item in tableC_analysis:
        competitor = item.get('Competitor', '')
        analysis = item.get('Analysis', '')

        # Try to parse the Analysis JSON; if it fails, skip this row
        try:
            analysis_data = json.loads(analysis.strip())
        except json.JSONDecodeError:
            print(f"Invalid JSON for analysis: {analysis}")
            continue
        
        # Add the competitor name as the column header
        competitors.append(competitor)
        
        # Flatten the analysis_data dictionary to get the key-value pairs as rows
        processed_data_tableC.append(analysis_data)

    # Convert the processed data into a DataFrame
    df = pd.DataFrame(processed_data_tableC)

    # Perform any necessary trimming or cleanup
    df = df.applymap(lambda x: x.strip() if isinstance(x, str) else x)

    # Add the "Parameters" column as the first column with the attribute names
    # The list of attribute names (keys) from the first competitor's analysis will be used as parameters
    parameters = list(df.columns)
    #df.insert(0, "Parameters", parameters)

    # Transpose the DataFrame so that competitors are in the headers
    transposed_df = df.T

    transposed_df.reset_index(inplace=True)

    transposed_df.rename(columns={'index': 'Parameters'}, inplace=True)

    competitors.insert(0,"Parameters")

    # Set the correct column headers (i.e., the competitor names)
    transposed_df.columns = competitors

    return transposed_df


def set_cell_background(cell, color):
    """Set cell background shading color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)  # Set fill color
    tcPr.append(shd)

def set_table_border(table):
    """Set table borders using XML manipulation."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        element = OxmlElement(f'w:{border}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        borders.append(element)
    tblPr.append(borders)


def create_new_doc():
    """Create a new Word document and return it."""
    doc = Document()
    doc.add_heading("Web Scraping Tables", level=1)
    return doc

def add_dataframe_to_doc(doc, df, table_number):
    """Add a DataFrame as a table to the Word document."""
    if df.empty:
        print(f"Table-{table_number}: The DataFrame is empty. No table will be created.")
        return
    
    # Add a bold title for the table
    table_extraction_paragraph = doc.add_paragraph(style='Normal')
    run = table_extraction_paragraph.add_run(f'Table-{table_number}')
    run.bold = True

    # Create a table in the document
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells

    # Add column headers
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)
        hdr_cells[j].paragraphs[0].runs[0].font.bold = True
        set_cell_background(hdr_cells[j], "F2F2F2")

    # Add rows from DataFrame
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, cell in enumerate(row):
            row_cells[j].text = str(cell)

    # Apply borders
    set_table_border(table)

def save_doc_to_file(doc, filename):
    """Save the Word document to a file."""
    doc.save(filename)

def Equivalence_Generation(row):
    prompt = f"""
     Analyze the following search row: {row} about a parameter, The first value contains the name of a parameter. The other values contain
     the information about that parameter for different device series, compare them among each other and give a one line comparison of at max 20 words. 
     Mention the similarities and differences just about the parameter dont include brand names or external information in your comparision. 
     Dont compare the Parameters that dont make sense to be compared in a device comparison table like License Holder, License Number, etc.
     Comparison should be apple to apple comparison, dont compare different traits between them. 
     There should be strictly no model response. 
    """

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an expert in competitor product analysis."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2

    )
    start_time = time.time()
    # Return the generated content if available
    if response and response['choices']:
        input_tokens = response['usage']['prompt_tokens']
        output_tokens = response['usage']['completion_tokens']
        total_tokens = response['usage']['total_tokens']
        all_tokens.append({
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': total_tokens,
            'response_time': time.time() - start_time
        })
    return response['choices'][0]['message']['content']

def extract_pdf_text(file_bytes, page_number=[]):
    """
    Extract text row-wise from the PDF, better handling of tables and structured content.
    """
    extracted_text = []

    # Open the PDF from bytes
    # doc = fitz.open(pdf_file)
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    if len(page_number) == 0:
        # Loop over the first 45 pages or up to total page count
        for page_num in range(min(45, doc.page_count)):
            try:
                page = doc.load_page(page_num)
                blocks = page.get_text("blocks")  # Extract text blocks with positions

                # Sort blocks by their vertical position (y-coordinate)
                blocks = sorted(blocks, key=lambda block: block[1])  # block[1] is the y-coordinate of the block

                current_row = []
                last_y_position = None
                y_threshold = 8  # Adjust threshold as needed to determine row breaks

                for block in blocks:
                    block_text = block[4]  # Text content of the block
                    y_position = block[1]  # y-coordinate of the block

                    # If this block's y-position is close to the last one, add it to the current row
                    if last_y_position is None or abs(y_position - last_y_position) < y_threshold:
                        current_row.append(block_text)
                    else:
                        # If the block is far from the last y-position, store the current row and start a new one
                        extracted_text.append(" ".join(current_row))
                        current_row = [block_text]  # Start a new row with the current block

                    last_y_position = y_position

                # Don't forget to add the last row if there's any content
                if current_row:
                    extracted_text.append(" ".join(current_row))

            except Exception as e:
                print(f"Failed to load page {page_num}: {e}")
    else:
        for page_num in page_number:
            try:
                page = doc.load_page(page_num)
                blocks = page.get_text("blocks")  # Extract text blocks with positions

                # Sort blocks by their vertical position (y-coordinate)
                blocks = sorted(blocks, key=lambda block: block[1])  # block[1] is the y-coordinate of the block

                current_row = []
                last_y_position = None
                y_threshold = 8  # Adjust threshold as needed to determine row breaks

                for block in blocks:
                    block_text = block[4]  # Text content of the block
                    y_position = block[1]  # y-coordinate of the block

                    # If this block's y-position is close to the last one, add it to the current row
                    if last_y_position is None or abs(y_position - last_y_position) < y_threshold:
                        current_row.append(block_text)
                    else:
                        # If the block is far from the last y-position, store the current row and start a new one
                        extracted_text.append(" ".join(current_row))
                        current_row = [block_text]  # Start a new row with the current block

                    last_y_position = y_position

                # Don't forget to add the last row if there's any content
                if current_row:
                    extracted_text.append(" ".join(current_row))

            except Exception as e:
                print(f"Failed to load page {page_num}: {e}")

    return extracted_text



def merge_cells_based_on_value(df, output_file, sheet_name="Sheet1"):
    # Create a new workbook
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_name

    # Write DataFrame to Excel sheet
    for row_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), start=1):
        for col_idx, value in enumerate(row, start=1):
            sheet.cell(row=row_idx, column=col_idx, value=value)

    # Find the column index of "comparison"
    comparison_column_idx = list(df.columns).index("comparison") + 1  # Excel is 1-based

    # Initialize variables for merging cells
    start_row = None
    current_value = None

    # Iterate through rows, starting from the second row (data rows)
    for row in range(2, len(df) + 2):  # +2 for header and 1-based index
        cell_value = sheet.cell(row=row, column=comparison_column_idx).value

        # Check if we need to start a new merge group
        if cell_value != current_value:
            # Merge the previous group
            if start_row is not None:
                sheet.merge_cells(
                    start_row=start_row,
                    start_column=comparison_column_idx,
                    end_row=row - 1,
                    end_column=comparison_column_idx
                )
                sheet.cell(start_row, comparison_column_idx).alignment = Alignment(horizontal="center", vertical="center")
            # Start a new group
            start_row = row
            current_value = cell_value

    # Merge the last group
    if start_row is not None:
        sheet.merge_cells(
            start_row=start_row,
            start_column=comparison_column_idx,
            end_row=len(df) + 1,
            end_column=comparison_column_idx
        )
        sheet.cell(start_row, comparison_column_idx).alignment = Alignment(horizontal="center", vertical="center")

    # Save the workbook
    workbook.save(output_file)
    print(f"File saved with merged cells to: {output_file}")

def scrapping(thermoDevice, previousGenDevice, user_keyword, refined_keyword, technical_data_sheet, prev_gen_data_sheet ):

    # thermoDevice_data_string = extract_pdf_text(technical_data_sheet)
    # previousGenDevice_data_string = extract_pdf_text(prev_gen_data_sheet)
    # Read and process technical data sheets for the current device
    thermoDevice_data_string = []
    for uploaded_file in technical_data_sheet:
        file_bytes = uploaded_file.read()  # Read the file content into bytes
        extracted_text = extract_pdf_text(file_bytes)
        thermoDevice_data_string.extend(extracted_text)

    # Read and process technical data sheets for the previous-gen device
    previousGenDevice_data_string = []
    for uploaded_file in prev_gen_data_sheet:
        file_bytes = uploaded_file.read()  # Read the file content into bytes
        extracted_text = extract_pdf_text(file_bytes)
        previousGenDevice_data_string.extend(extracted_text)
    # Initialize WebDriver
    driver = initialize_driver()

    try:

        # Analyze data
        TableA_analysis = analyze_google_data_tableA(
            thermoDevice, previousGenDevice, thermoDevice_data_string, previousGenDevice_data_string
        )

        # Create DataFrame
        tableA_df = create_dataframe_from_analysis(TableA_analysis, previousGenDevice, thermoDevice)
        tableA_df = tableA_df.sort_values(by=tableA_df.columns[-1], ascending=False)
        print(tableA_df)
        # Usage
        # output_file_path = "Features/tableA_with_merged.xlsx"
        # merge_cells_based_on_value(tableA_df, output_file_path)
        # tableA_df.to_excel("Features/tableA.xlsx", index=False)

        # # Scrape data
        # thermoDevice_data = scrape_device_details(driver, thermoDevice)
        # previousGenDevice_data = scrape_device_details(driver, previousGenDevice)

        # # Format scraped data
        # thermoDevice_data_string = format_scraped_data(thermoDevice_data)
        # previousGenDevice_data_string = format_scraped_data(previousGenDevice_data)

        # # Analyze data
        # TableA_analysis = analyze_google_data_tableA(
        #     thermoDevice, previousGenDevice, thermoDevice_data_string, previousGenDevice_data_string
        # )

        # Create DataFrame
        # tableA_df = create_dataframe_from_analysis(TableA_analysis, previousGenDevice, thermoDevice)
        print("tableA_df created sucesssfully")
        
        '''
        Below commeneted code need to update

        '''
        # scrap_process()
        # print("scrap process completed")
        file_path = r"Features/MD - Medical Devices (3).xlsx"
        sheet_name = "Sheet1"
        tableB_df = scrap_process_tableB_df(file_path, sheet_name, thermoDevice, previousGenDevice, user_keyword, refined_keyword)
        columnHeaders=[]
        for item in competitor_analysis:
            analysis = json.loads(item['Analysis'])  # Parse the JSON string into a dictionary
            company_name = analysis['Company Name']  # Extract the company name
            
            # Step 2: Compare it with the 'Company Name' column in df
            if company_name in tableB_df['Company Name'].values:
                columnHeaders.append(item['Competitor']) 
        tableB_forC=tableB_df
        tableB_forC["Column_Header"]=columnHeaders

        base_description = tableB_forC.iloc[0]['Description']
        # Step 2: Extract the other descriptions
        other_descriptions = tableB_forC.iloc[1:]['Description']
        # Step 3: Calculate TF-IDF vectors for the descriptions
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform([base_description] + other_descriptions.tolist())
        # Step 4: Compute cosine similarity between the first row (base description) and the other descriptions
        cosine_similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
        # Step 5: Get the indices of the top 2 most similar descriptions
        top_6_indices = cosine_similarities.argsort()[-6:][::-1]
        # Step 6: Store the first row and the top 2 matches in a new DataFrame
        top_6_matches = tableB_forC.iloc[top_6_indices + 1]  # +1 because we excluded the first row
        tableC_input = pd.concat([tableB_forC.iloc[[0]], top_6_matches])

        # Step 1: Extract the 'Company Name' from the 'Analysis' field in demo_list
        competitor_matches_C_input = []
        competitor_matches_C_input_for_search = []
        
        for item in competitor_analysis:
            analysis = json.loads(item['Analysis'])  # Parse the JSON string into a dictionary
            company_name = analysis['Company Name']  # Extract the company name
            
            # Step 2: Compare it with the 'Company Name' column in tableC_input
            if company_name in tableC_input['Company Name'].values:
                competitor_matches_C_input.append(item['Competitor'])
                
                # Check if the competitor is present in the "Column_Header" column
                matching_row = tableC_input[tableC_input['Column_Header'] == item['Competitor']]
                if not matching_row.empty:
                    # Concatenate "Company Name" and "Product Name" and append to competitor_matches_C_input_for_search
                    concatenated_value = matching_row.iloc[0]['Company Name'] + " " + matching_row.iloc[0]['Product Name']
                    competitor_matches_C_input_for_search.append(concatenated_value)
                

        competitor_matches_C_input_for_search[0]=thermoDevice
        # driver.get("https://www.google.com")
        tableC_analysis=[]
        for search, competitor in zip(competitor_matches_C_input_for_search,competitor_matches_C_input):
            try:
                print(f"Searching for: {search}")
                google_data = scrape_device_details(driver, search)
                google_data_string = "\n".join(
                    f"Result Number: {item['Result Number']}\n"
                    f"URL: {item['URL']}\n"
                    f"Title: {item['Title']}\n"
                    f"Content: {item['Content']}\n"
                    for item in google_data
                )
                analysis = analyze_google_data_TableC(google_data_string, competitor)
                print(analysis)
                tableC_analysis.append({"Competitor": competitor, "Analysis": analysis})
            except Exception as e:
                print(f"Error during analysis for competitor '{competitor}': {e}")

        print("---------------------------------------------------------")
        ##print(tableC_df)
        print(tableC_analysis)

        # Process and transpose the tableC_analysis data
        df = process_and_transpose_data(tableC_analysis)

        # Display the transposed DataFrame with Parameters in the top-left corner
        print("---------------------------------------------------------")
        print(df)

        # df = tableC_df
        df1 = pd.read_excel(file_path)

        df1['concatenated_column'] = df1['Name of Device and Device Class'].astype(str) + ' ' + df1['Brand Name'].astype(str)

        # Step 2: Filter rows based on the column names in `df`
        filter_values = df.columns[2:].tolist()  # Use column names from df for filtering
        filtered_df1 = df1[df1['concatenated_column'].isin(filter_values)]
        
        print("------------------------------------------")
        print(filtered_df1)

        filtered_df1['License Number'] = filtered_df1['License Number'].str.extract(r'(\S{3}/\S+)')

        # Step 4: Create the final DataFrame with only 'Device and Brand Name' and 'License Number'
        final_df = filtered_df1[['concatenated_column', 'License Number']]

        # Step 5: Ensure all filter values are accounted for in `final_df`
        for value in filter_values:
            if value not in final_df['concatenated_column'].values:
                # Append value and set License Number as null
                new_row = pd.DataFrame({
                    'concatenated_column': [value],
                    'License Number': [None]
                })
                final_df = pd.concat([final_df, new_row], ignore_index=True)

        # Step 6: Clean up and rename columns
        final_df.reset_index(drop=True, inplace=True)
        final_df.index += 1  # Adjust the index to start from 1
        final_df.rename(columns={'concatenated_column': 'Device and Brand name'}, inplace=True)

        # Step 7: Map License Numbers to the Original DataFrame
        license_number_row = []
        for device in df.columns[1:]:  # Skip the 'Parameters' column
            matching_row = final_df[final_df['Device and Brand name'] == device]
            if not matching_row.empty:
                # Append the corresponding license number
                license_number_row.append(matching_row['License Number'].values[0])
            else:
                # If no match, append None
                license_number_row.append(None)

        # Step 8: Add the new row with License Numbers to the original DataFrame
        license_number_row = ['License Number'] + license_number_row
        new_row_df = pd.DataFrame([license_number_row], columns=df.columns)

        # Step 3: Concatenate original DataFrame with the new row
        final_df = pd.concat([df.iloc[:1], new_row_df, df.iloc[1:]], ignore_index=True)
        
        
        equivalence=[]
        for row in final_df.values:
            print(row)
            equivalence_string = Equivalence_Generation(row)
            equivalence.append(equivalence_string)

        final_df['Equivalence']=equivalence
        tableC_df = final_df
        print(tableC_df)
        if 'Column_Header' in tableB_df.columns:
            tableB_df = tableB_df.drop(columns=['Column_Header'])
        doc = create_new_doc()  # Calls the `create_new_doc` function to initialize a new Word document

        # Step 3: Add DataFrames one by one
        table_counter = 1  # Initialize the table counter

        # Add the first DataFrame
        add_dataframe_to_doc(doc, tableA_df, table_counter)  # Calls `add_dataframe_to_doc` to add df1 to the document
        table_counter += 1  # Increment the table counter

        add_dataframe_to_doc(doc, tableB_df, table_counter)  # Calls `add_dataframe_to_doc` to add df2 to the document
        table_counter += 1  # Increment the table counter

        # Add the second DataFrame
        add_dataframe_to_doc(doc, tableC_df, table_counter)  # Calls `add_dataframe_to_doc` to add df2 to the document
        # save_doc_to_file(doc, "web_scrapping_tables_final.docx")  
        # Save the document to a byte stream
        byte_stream = io.BytesIO()
        doc.save(byte_stream)
        byte_stream.seek(0)
        
        # Step 4: Save the document
        
        print(all_tokens)
        # Calculate the sum of input_tokens, output_tokens, and total_tokens
        token_info = {
            "input_tokens": sum(entry['input_tokens'] for entry in all_tokens),
            "output_tokens": sum(entry['output_tokens'] for entry in all_tokens),
            "total_tokens": sum(entry['total_tokens'] for entry in all_tokens),
        }
        # # Remove all .xlsx files in Features folder
        # for file in os.listdir('Features'):
        #     if file.endswith('.xlsx'):
        #         os.remove(os.path.join('Features', file))
        print("success")   
        return byte_stream, token_info     

    finally:
        # Close the driver
        driver.quit()
