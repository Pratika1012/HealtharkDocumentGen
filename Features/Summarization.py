# new(12/11/2024)

from langchain import PromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.chains import LLMChain
from langchain.chains.summarize import load_summarize_chain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document  # For reading and writing Word files
import logging
import tiktoken
import re
import openai
import json
import time
import os
import pdfplumber
from io import BytesIO
from doc_generate import save_summaries_to_word
import streamlit as st

with open('Config/configuration.json', 'r') as f:
    config = json.load(f)

openai.api_key = config['api_key']  # Set the API key from the configuration file

def count_tokens(text, model="gpt-4o"):
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))

# Initialize the OpenAI model
llm = ChatOpenAI(
    temperature=config['generation_config']['temperature'],
    model_name=config['model_name'],
    openai_api_key=config['api_key']
)

def extract_text_from_pdf(file):
    """
    Extracts text from a PDF file using pdfplumber.
    """
    try:
        with pdfplumber.open(file) as pdf:
            all_text = ""
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text += text
            return all_text
    except FileNotFoundError:
        logging.error(f"File not found: {file}")
        return ""
    except pdfplumber.PDFSyntaxError as e:
        logging.error(f"PDF syntax error in file {file}: {e}")
        return ""
    except Exception as e:
        logging.error(f"Unexpected error while extracting text from PDF: {e}")
        return ""

def read_word_file(file_path):
    """
    Reads a Word document and extracts text from its paragraphs.
    """
    try:
        doc = Document(file_path)
        text = [para.text for para in doc.paragraphs if para.text]
        return '\n'.join(text)
    except FileNotFoundError:
        logging.error(f"File not found: {file_path}")
        return ""
    except Exception as e:
        logging.error(f"Error reading Word file {file_path}: {e}")
        return ""

def extract_instructions_from_template(template_path):
    """
    Extracts instructions from a Word document template.
    """
    try:
        doc = Document(template_path)
        instructions = "\n".join([para.text for para in doc.paragraphs])
        return instructions
    except FileNotFoundError:
        logging.error(f"Template file not found: {template_path}")
        return ""
    except Exception as e:
        logging.error(f"Error reading template file {template_path}: {e}")
        return ""

def summarization(file_stream, file_type, instructions):

    token_info = {}

    if file_type == 'docx':
        text = read_word_file(file_stream)
    elif file_type == 'pdf':
        text = extract_text_from_pdf(file_stream)

    text_token_count = count_tokens(text)
    token_info['text_tokens'] = text_token_count

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=100)
    text_chunks = text_splitter.create_documents([text])
    logging.info(f'Created {len(text_chunks)} chunks from the text.')

    summary_prompt = f"""
    Follow these instructions to summarize the document:
    {instructions}
    {{text}}
    """

    if len(text_chunks) > 1:
        chunks_prompt = """
            Please summarize the below text:
            Text: `{text}`
            Summary:
        """
        map_prompt_template = PromptTemplate(input_variables=['text'], template=chunks_prompt)
        summary_prompt_template = PromptTemplate(input_variables=['text'], template=summary_prompt)

        summary_chain = load_summarize_chain(
            llm=llm,
            chain_type='map_reduce',
            map_prompt=map_prompt_template,
            combine_prompt=summary_prompt_template,
            verbose=True
        )
    else:
        summary_prompt_template = PromptTemplate(input_variables=['text'], template=summary_prompt)
        summary_chain = load_summarize_chain(
            llm=llm,
            chain_type='stuff',
            prompt=summary_prompt_template,
            verbose=True
        )

    summary = summary_chain.run(input_documents=text_chunks)
    summary_token_count = count_tokens(summary)
    logging.info(f'Summary token count: {summary_token_count}')
    token_info['summary_tokens'] = summary_token_count

    return summary, token_info

def combine_summaries_with_template(all_summaries, template):
    combined_text = ""
    for file_name, summary in all_summaries:
        combined_text += f"Summary of {file_name}:\n{summary}\n\n"

    prompt_template = f"""
    Given the following template and Summarized text, use the template only to position each section of the Summarized text
    according to the order in the template. **Dont change, rephrase, concise any part of the content.**
    Ensure the summaries are placed correctly and make any necessary tweaks only if any mentioned in template for that section:

    Template: {template}

    Summarized text:
    {combined_text}

    Once ordered make all the headings and sub-headings defined in the template bold by encapsulating them between two asterisk eg. **this text should be bold**
    **Note:-Any Sub-sections defined in the template should be bulletted by alphabets like a,b,c.**
    """

    prompt = PromptTemplate(
        input_variables=["template", "combined_text"],
        template=prompt_template
    )

    llm_chain = LLMChain(llm=llm, prompt=prompt, verbose=True)
    final_document = llm_chain.run(template=template, combined_text=combined_text)

    return final_document

# MAIN Function
def summarized_Document(input_files, template_path):
    instructions = extract_instructions_from_template(template_path)
    if input_files:
        all_summaries = []
        token_info_list = []
        

        for document in input_files:
            file_stream = BytesIO(document.read())
            file_type = document.name.split('.')[-1]
            summary, token_info = summarization(file_stream, file_type, instructions)
                
            all_summaries.append((document.name, summary))
            token_info_list.append(token_info)

        final_output = combine_summaries_with_template(all_summaries, instructions)

        input_tokens = sum(info.get('text_tokens', 0)  for info in token_info_list)
        output_tokens = sum( info.get('summary_tokens', 0) for info in token_info_list)
        total_tokens = sum(info.get('text_tokens', 0) + info.get('summary_tokens', 0) for info in token_info_list)

        
        if all_summaries:
            summary_doc = save_summaries_to_word(final_output)
            logging.info("process_Summary_text_with_GPT...............")
            logging.info(f"Input Tokens: {input_tokens}")
            logging.info(f"Output Tokens: {output_tokens}")
            logging.info(f"Total Tokens: {total_tokens}")
            
            return summary_doc,  {
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': total_tokens
        }
        else:
            return None,None

